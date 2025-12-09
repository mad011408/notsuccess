"""NEXUS AI Agent - Document Loader"""

import os
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, field
from pathlib import Path

from config.logging_config import get_logger


logger = get_logger(__name__)


@dataclass
class Document:
    """Loaded document"""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    source: str = ""
    doc_type: str = ""


class DocumentLoader:
    """Load documents from various sources"""

    def __init__(self):
        self._loaders = {
            '.txt': self._load_text,
            '.md': self._load_text,
            '.py': self._load_text,
            '.json': self._load_json,
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.csv': self._load_csv,
            '.html': self._load_html,
        }

    def load(self, path: str) -> Document:
        """
        Load a document from file

        Args:
            path: File path

        Returns:
            Document object
        """
        path = os.path.abspath(path)

        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found: {path}")

        ext = os.path.splitext(path)[1].lower()
        loader = self._loaders.get(ext, self._load_text)

        content = loader(path)

        return Document(
            content=content,
            metadata={
                "filename": os.path.basename(path),
                "extension": ext,
                "size": os.path.getsize(path),
                "modified": os.path.getmtime(path)
            },
            source=path,
            doc_type=ext[1:] if ext else "unknown"
        )

    def load_directory(
        self,
        path: str,
        extensions: Optional[List[str]] = None,
        recursive: bool = True
    ) -> List[Document]:
        """
        Load all documents from directory

        Args:
            path: Directory path
            extensions: Filter by extensions
            recursive: Search recursively

        Returns:
            List of Document objects
        """
        documents = []
        path = os.path.abspath(path)

        if not os.path.isdir(path):
            raise NotADirectoryError(f"Not a directory: {path}")

        if recursive:
            for root, _, files in os.walk(path):
                for filename in files:
                    file_path = os.path.join(root, filename)
                    ext = os.path.splitext(filename)[1].lower()

                    if extensions and ext not in extensions:
                        continue

                    try:
                        doc = self.load(file_path)
                        documents.append(doc)
                    except Exception as e:
                        logger.warning(f"Failed to load {file_path}: {e}")
        else:
            for filename in os.listdir(path):
                file_path = os.path.join(path, filename)
                if not os.path.isfile(file_path):
                    continue

                ext = os.path.splitext(filename)[1].lower()
                if extensions and ext not in extensions:
                    continue

                try:
                    doc = self.load(file_path)
                    documents.append(doc)
                except Exception as e:
                    logger.warning(f"Failed to load {file_path}: {e}")

        return documents

    def _load_text(self, path: str) -> str:
        """Load text file"""
        encodings = ['utf-8', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        with open(path, 'rb') as f:
            return f.read().decode('utf-8', errors='replace')

    def _load_json(self, path: str) -> str:
        """Load JSON file"""
        import json
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return json.dumps(data, indent=2)

    def _load_pdf(self, path: str) -> str:
        """Load PDF file"""
        try:
            import PyPDF2
            text_parts = []
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("PyPDF2 not installed. PDF loading unavailable.")
            return ""

    def _load_docx(self, path: str) -> str:
        """Load DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            logger.warning("python-docx not installed. DOCX loading unavailable.")
            return ""

    def _load_csv(self, path: str) -> str:
        """Load CSV file"""
        import csv
        rows = []
        with open(path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(", ".join(row))
        return "\n".join(rows)

    def _load_html(self, path: str) -> str:
        """Load HTML file"""
        try:
            from bs4 import BeautifulSoup
            with open(path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            return soup.get_text(separator='\n', strip=True)
        except ImportError:
            # Fallback to regex
            import re
            with open(path, 'r', encoding='utf-8') as f:
                html = f.read()
            text = re.sub(r'<[^>]+>', '', html)
            return text

    def load_url(self, url: str) -> Document:
        """Load document from URL"""
        import urllib.request

        with urllib.request.urlopen(url) as response:
            content = response.read().decode('utf-8', errors='replace')

        return Document(
            content=content,
            metadata={"url": url},
            source=url,
            doc_type="url"
        )

    def add_loader(self, extension: str, loader_func) -> None:
        """Add custom loader for extension"""
        self._loaders[extension] = loader_func

    def supported_extensions(self) -> List[str]:
        """Get list of supported extensions"""
        return list(self._loaders.keys())

