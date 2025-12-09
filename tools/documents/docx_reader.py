"""NEXUS AI Agent - DOCX Reader"""

from typing import Optional, List, Dict, Any
from dataclasses import dataclass, field


@dataclass
class DocxParagraph:
    """DOCX paragraph"""
    text: str
    style: str = ""
    is_heading: bool = False
    heading_level: int = 0


@dataclass
class DocxTable:
    """DOCX table"""
    rows: List[List[str]]


@dataclass
class DocxDocument:
    """DOCX document content"""
    path: str
    paragraphs: List[DocxParagraph] = field(default_factory=list)
    tables: List[DocxTable] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None


class DocxReader:
    """Read DOCX documents"""

    def read(self, path: str) -> DocxDocument:
        """
        Read DOCX document

        Args:
            path: DOCX file path

        Returns:
            DocxDocument object
        """
        try:
            from docx import Document

            doc = DocxDocument(path=path)
            docx = Document(path)

            # Extract metadata
            core_props = docx.core_properties
            doc.metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            }

            # Extract paragraphs
            for para in docx.paragraphs:
                style_name = para.style.name if para.style else ""
                is_heading = style_name.startswith('Heading')
                heading_level = 0

                if is_heading:
                    try:
                        heading_level = int(style_name.replace('Heading ', ''))
                    except ValueError:
                        heading_level = 1

                doc.paragraphs.append(DocxParagraph(
                    text=para.text,
                    style=style_name,
                    is_heading=is_heading,
                    heading_level=heading_level
                ))

            # Extract tables
            for table in docx.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    rows.append(cells)
                doc.tables.append(DocxTable(rows=rows))

            return doc

        except ImportError:
            return DocxDocument(
                path=path,
                error="python-docx not installed. Run: pip install python-docx"
            )
        except Exception as e:
            return DocxDocument(path=path, error=str(e))

    def get_text(self, path: str) -> str:
        """Get all text from DOCX"""
        doc = self.read(path)
        if doc.error:
            return ""
        return "\n".join(p.text for p in doc.paragraphs)

    def get_headings(self, path: str) -> List[Dict[str, Any]]:
        """Get all headings"""
        doc = self.read(path)
        if doc.error:
            return []

        return [
            {"text": p.text, "level": p.heading_level}
            for p in doc.paragraphs
            if p.is_heading
        ]

    def get_tables(self, path: str) -> List[List[List[str]]]:
        """Get all tables"""
        doc = self.read(path)
        if doc.error:
            return []
        return [t.rows for t in doc.tables]

    def get_metadata(self, path: str) -> Dict[str, Any]:
        """Get document metadata"""
        doc = self.read(path)
        return doc.metadata

    def search(self, path: str, query: str, case_sensitive: bool = False) -> List[Dict[str, Any]]:
        """Search in document"""
        doc = self.read(path)
        if doc.error:
            return []

        results = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text if case_sensitive else para.text.lower()
            search_query = query if case_sensitive else query.lower()

            if search_query in text:
                results.append({
                    "paragraph_index": i,
                    "text": para.text,
                    "style": para.style
                })

        return results

    def to_markdown(self, path: str) -> str:
        """Convert DOCX to Markdown"""
        doc = self.read(path)
        if doc.error:
            return ""

        md_parts = []
        for para in doc.paragraphs:
            if para.is_heading and para.heading_level > 0:
                prefix = "#" * para.heading_level
                md_parts.append(f"{prefix} {para.text}")
            elif para.text.strip():
                md_parts.append(para.text)
            else:
                md_parts.append("")

        return "\n\n".join(md_parts)

